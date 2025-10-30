import streamlit as st
from groq import Groq
from io import BytesIO
import json, re
from docx import Document

st.set_page_config(page_title="Lexi — Legal Document Filler", layout="wide")

st.title("Lexi — Your AI Legal Document Assistant (Groq Llama 3.3 70B)")
st.write(
    "Lexi helps you complete legal documents intelligently. Upload a `.docx` file with "
    "placeholders like [Company Name], [Effective Date], or $[Amount], and Lexi will "
    "analyze each placeholder, explain what it means, and chat with you to fill it out "
    "step by step.  You can generate the final document only when every field is filled."
)

# ---------------- Sidebar ----------------
groq_api_key = st.sidebar.text_input(
    "Groq API Key", type="password", key="groq_api_key_sidebar"
)
if not groq_api_key:
    st.warning("Please enter your Groq API key to start Lexi.")
    st.stop()
client = Groq(api_key=groq_api_key)

# ---------------- Helpers ----------------
def load_docx(data): return Document(data)

def extract_placeholders_with_context(doc, window=1):
    regex = re.compile(r"(\$?\[+[^\[\]]+\]+)")
    placeholders, contexts = [], {}
    for i,p in enumerate(doc.paragraphs):
        found = regex.findall(p.text)
        if found:
            context=[]
            for off in range(-window,window+1):
                j=i+off
                if 0<=j<len(doc.paragraphs):
                    t=doc.paragraphs[j].text.strip()
                    if t: context.append(t)
            for f in found:
                placeholders.append(f)
                contexts[f]=" ".join(context)
    uniq, seen = [], set()
    for ph in placeholders:
        if ph not in seen:
            uniq.append(ph)
            seen.add(ph)
    return uniq, contexts

def replace_placeholders(doc,mapping):
    def repl(txt):
        for k,v in mapping.items(): txt=txt.replace(k,v)
        return txt
    for p in doc.paragraphs:
        if any(k in p.text for k in mapping):
            new=repl(p.text)
            for i in range(len(p.runs)-1,-1,-1):
                p.runs[i]._element.getparent().remove(p.runs[i]._element)
            p.add_run(new)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if any(k in c.text for k in mapping):
                    c._tc.clear_content()
                    c.add_paragraph(repl(c.text))
    return doc

def get_missing(values): return [k for k,v in values.items() if not v.strip()]

def analyze_placeholder_contexts(placeholders,contexts,client):
    results={}
    for ph in placeholders:
        ctx=contexts.get(ph,"")
        prompt=f"""
You are a legal-document analyst.
Describe what information this placeholder expects and give one realistic example.
Respond only in JSON:
{{"description":"...","example":"..."}}

Placeholder: {ph}
Context: {ctx}
"""
        try:
            r=client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role":"user","content":prompt}],
                temperature=0.2,
                max_tokens=200,
            )
            text=r.choices[0].message.content
            m=re.search(r"\{.*\}",text,re.S)
            results[ph]=json.loads(m.group(0)) if m else {"description":text.strip(),"example":""}
        except Exception as e:
            results[ph]={"description":f"Analysis failed: {e}","example":""}
    return results

# ---------------- Upload ----------------
uploaded = st.file_uploader("Upload your .docx document", type=["docx"], key="main_docx_uploader")
if not uploaded:
    st.info("Please upload a .docx file to continue.")
    st.stop()

doc = load_docx(uploaded)
placeholders, contexts = extract_placeholders_with_context(doc)
if not placeholders:
    st.warning("No placeholders found – ensure they use [brackets].")
    st.stop()

st.subheader(f"Detected {len(placeholders)} placeholders")

# ---------------- AI Analysis ----------------
st.write("Analyzing placeholder contexts with Groq Llama 3.3 70B …")
analysis = analyze_placeholder_contexts(placeholders, contexts, client)

st.subheader("AI-Generated Insights (Lexi’s Understanding)")
for ph in placeholders:
    info = analysis.get(ph, {})
    with st.expander(ph):
        st.markdown(f"**Context:** {contexts.get(ph,'')}")
        st.markdown(f"**Lexi’s Insight:** {info.get('description','')}")
        if info.get("example"):
            st.markdown(f"**Example:** {info['example']}")

# ---------------- Session Init ----------------
if "chat_history" not in st.session_state:
    st.session_state.chat_history=[]
if "placeholder_values" not in st.session_state:
    st.session_state.placeholder_values={ph:"" for ph in placeholders}

# --------------- Lexi’s Personality ---------------
def lexi_intro():
    intro_msg = (
        "Hi there! I’m **Lexi**, your AI legal assistant. "
        "I’ve reviewed your document and found several placeholders that need to be filled. "
        "My job is to guide you through each one, making sure everything is clear and complete. "
        "Let’s start! You can tell me information like ‘The company name is Acme Corp’ or ‘Effective Date is November 1, 2025’. "
        "Once we’ve filled everything, I’ll generate your final document."
    )
    return intro_msg

# Show Lexi’s intro once after upload
if not st.session_state.chat_history:
    st.session_state.chat_history.append({"role":"assistant","content":lexi_intro()})

# ---------------- Chat Logic ----------------
def groq_conversation(user_msg,history,placeholders,values,contexts,analysis):
    missing=get_missing(values)
    meanings="\n".join(f"{ph}: {analysis.get(ph,{}).get('description','')}" for ph in placeholders)
    context_text="\n".join(f"{ph}: {contexts.get(ph,'')}" for ph in placeholders)

    sys_prompt = (
        "You are Lexi, a friendly but professional AI legal assistant helping the user fill placeholders in a legal document. "
        "Be concise, courteous, and encouraging. "
        "Use the semantic hints and context to interpret user inputs. "
        "Always include a JSON mapping of placeholder → value after your message. "
        "If any placeholders remain empty, gently remind the user which ones are still missing."
    )

    msgs=[{"role":"system","content":sys_prompt}]+history
    msgs.append({
        "role":"user",
        "content":(
            f"Current mapping: {json.dumps(values)}\n"
            f"Missing: {missing}\n"
            f"Placeholder meanings:\n{meanings}\n"
            f"Contexts:\n{context_text}\n\n"
            f"User message: {user_msg}"
        ),
    })

    resp=client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=msgs,
        temperature=0.4,
        max_tokens=800,
    )
    content=resp.choices[0].message.content
    reply=content
    m=re.search(r"\{.*\}",content,re.S)
    if m:
        try:
            parsed=json.loads(m.group(0))
            for k,v in parsed.items():
                for ph in placeholders:
                    if k.strip().lower().replace(" ","") in ph.lower().replace(" ",""):
                        values[ph]=v
        except Exception:
            pass
    miss_after=get_missing(values)
    if miss_after:
        reply += f"\n\n(Remaining fields: {', '.join(miss_after)})"
    else:
        reply += "\n\n✅ Great! All placeholders are filled — I can generate the final document when you're ready."
    return reply, values

# ---------------- Chat UI ----------------
st.header("Chat with Lexi")

user_msg = st.chat_input("Type your response here …")

if user_msg:
    st.session_state.chat_history.append({"role":"user","content":user_msg})
    reply,updated=groq_conversation(
        user_msg,
        st.session_state.chat_history,
        placeholders,
        st.session_state.placeholder_values,
        contexts,
        analysis,
    )
    st.session_state.placeholder_values=updated
    st.session_state.chat_history.append({"role":"assistant","content":reply})

# Render chat
for msg in st.session_state.chat_history:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

filled = sum(1 for v in st.session_state.placeholder_values.values() if v)
st.progress(filled/len(placeholders))
st.write(f"Filled {filled}/{len(placeholders)} placeholders.")
with st.expander("Current Placeholder Values"):
    st.json(st.session_state.placeholder_values)

# ---------------- Final Document (Gated) ----------------
# missing_now = get_missing(st.session_state.placeholder_values)
# if not missing_now:
#     st.success("✅ All placeholders are filled — Lexi can now generate the final document.")
#     if st.button("Generate Final Document"):
#         filled_doc=replace_placeholders(load_docx(uploaded),st.session_state.placeholder_values)
#         bio=BytesIO()
#         filled_doc.save(bio); bio.seek(0)
#         st.download_button("Download Completed Document", data=bio, file_name="filled_document.docx")
#         preview=[p.text for p in filled_doc.paragraphs[:30] if p.text.strip()]
#         st.subheader("Preview (first few paragraphs)")
#         st.text("\n\n".join(preview))
# else:
#     st.info(f"Lexi: We’re almost done — still need {', '.join(missing_now)} before generating the document.")

missing_now = get_missing(st.session_state.placeholder_values)

@st.dialog("Final Document Preview", width="large")
def show_final_dialog(filled_doc):
    st.write("Here’s your completed document. You can review and download it below.")

    bio = BytesIO()
    filled_doc.save(bio)
    bio.seek(0)

    st.download_button(
        "Download Completed Document",
        data=bio,
        file_name="filled_document.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    preview = [p.text for p in filled_doc.paragraphs[:30] if p.text.strip()]
    st.subheader("Preview (first few paragraphs)")
    st.text("\n\n".join(preview))
    st.info("Close this dialog to continue chatting with Lexi.")

# Show dialog only if all placeholders filled
if not missing_now:
    st.success("✅ All placeholders are filled — Lexi can now generate the final document.")
    if st.button("Generate Final Document"):
        filled_doc = replace_placeholders(load_docx(uploaded), st.session_state.placeholder_values)
        show_final_dialog(filled_doc)
else:
    st.info(f"Lexi: We’re almost done — still need {', '.join(missing_now)} before generating the document.")